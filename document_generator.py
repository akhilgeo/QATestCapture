from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet
import os

class DocumentGenerator:
    def generate_docx(self, data, output_filename="QA_Report.docx"):
        doc = Document()
        doc.add_heading('QA Test Report', 0)

        for item in data:
            doc.add_heading(f"Timestamp: {item['timestamp']}", level=2)
            doc.add_paragraph(f"Description: {item['description']}")
            doc.add_paragraph(f"Active Window: {item['window_title']}")
            
            if os.path.exists(item['image_path']):
                try:
                    doc.add_picture(item['image_path'], width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"[Error adding image: {e}]")
            
            doc.add_paragraph("-" * 50)

        doc.save(output_filename)
        return output_filename

    def generate_pdf(self, data, output_filename="QA_Report.pdf"):
        doc = SimpleDocTemplate(output_filename, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph("QA Test Report", styles['Title']))
        story.append(Spacer(1, 12))

        for item in data:
            story.append(Paragraph(f"Timestamp: {item['timestamp']}", styles['Heading2']))
            story.append(Paragraph(f"Description: {item['description']}", styles['Normal']))
            story.append(Paragraph(f"Active Window: {item['window_title']}", styles['Normal']))
            story.append(Spacer(1, 6))

            if os.path.exists(item['image_path']):
                try:
                    # Resize image to fit page if necessary (approx 6 inches width)
                    img = RLImage(item['image_path'], width=400, height=225) # Aspect ratio might be off, but keeps it simple
                    # Better to preserve aspect ratio
                    # img = RLImage(item['image_path'])
                    # img._restrictSize(400, 300)
                    story.append(img)
                except Exception as e:
                    story.append(Paragraph(f"[Error adding image: {e}]", styles['Normal']))
            
            story.append(Spacer(1, 12))
            story.append(Paragraph("-" * 50, styles['Normal']))
            story.append(Spacer(1, 12))

        doc.build(story)
        return output_filename
